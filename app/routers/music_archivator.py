from __future__ import annotations

import os
import io
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Request, UploadFile, File, Form
from sqlalchemy.orm import Session, joinedload
from app.db import get_db
from app.models import Instrument, Song, SongFile
from app.schemas import (
    InstrumentOut,
    InstrumentSetup,
    InstrumentCreate,
    SongOut,
    SongCreate,
    SongUpdate,
    SongFileUpload,
)
from app.token_store import TOKENS
from app.google_client import (
    drive_service,
    get_credentials_from_session,
    ensure_folder,
    upload_file_to_drive,
    update_or_create_file,
    _escape_q,
)
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_TAB_ALIGNMENT

router = APIRouter(prefix="/ma", tags=["MusicArchivator"])


def get_user_email(request: Request) -> str:
    sid = request.session.get("sid")
    if not sid or sid not in TOKENS:
        raise HTTPException(status_code=401, detail="Not logged in")
    user = TOKENS[sid].get("user")
    if not user or "email" not in user:
        raise HTTPException(status_code=401, detail="User email not found in session")
    return user["email"]


@router.get("/instruments", response_model=list[InstrumentOut])
def list_instruments(request: Request, db: Session = Depends(get_db)):
    email = get_user_email(request)
    return db.query(Instrument).filter(Instrument.owner_email == email).all()


@router.post("/instruments/setup")
def save_instrument_setup(
    payload: InstrumentSetup, request: Request, db: Session = Depends(get_db)
):
    email = get_user_email(request)

    # Simple strategy: delete old and insert new
    db.query(Instrument).filter(Instrument.owner_email == email).delete()

    for inst in payload.instruments:
        new_inst = Instrument(
            owner_email=email,
            name=inst.name,
            category=inst.category,
            is_tracked=inst.is_tracked,
        )
        db.add(new_inst)
    db.commit()

    # Aktualizujeme dokumenty (seznamy)
    try:
        sid = request.session.get("sid")
        token_data = TOKENS[sid]
        creds = get_credentials_from_session(token_data["token"])
        drv = drive_service(creds)
        sync_ma_documents(email, drv, db)
    except Exception as e:
        print(f"MA Sync Docs Error: {e}")

    return {"status": "ok"}


@router.get("/songs", response_model=list[SongOut])
def list_songs(request: Request, db: Session = Depends(get_db)):
    email = get_user_email(request)
    return (
        db.query(Song)
        .options(joinedload(Song.files))
        .filter(Song.owner_email == email)
        .all()
    )


@router.post("/songs", response_model=SongOut)
def create_song(payload: SongCreate, request: Request, db: Session = Depends(get_db)):
    try:
        email = get_user_email(request)

        # Validace čísla
        num = payload.number.strip()
        if num.upper() == "N":
            num = "N"
        elif not num.isdigit():
            raise HTTPException(
                status_code=400, detail="Číslo písně musí být celé číslo nebo 'N'"
            )

        if num != "N":
            existing = (
                db.query(Song)
                .filter(Song.owner_email == email, Song.number == num)
                .first()
            )
            if existing:
                raise HTTPException(
                    status_code=400, detail=f"Skladba s číslem {num} již existuje"
                )

        # Google Drive Logika
        sid = request.session.get("sid")
        token_data = TOKENS[sid]
        creds = get_credentials_from_session(token_data["token"])
        drv = drive_service(creds)

        root_id = os.getenv("BAND_DRIVE_ROOT_FOLDER_ID")
        # Přejmenováno na kratší "Noty"
        noty_folder_id = ensure_folder(drv, root_id, "Noty")

        song_folder_name = f"{num} {payload.title}"
        song_folder_id = ensure_folder(drv, noty_folder_id, song_folder_name)

        new_song = Song(
            owner_email=email,
            number=num,
            title=payload.title,
            singer=payload.singer,
            duration=payload.duration,
            category=payload.category,
            drive_folder_id=song_folder_id,
        )
        db.add(new_song)
        db.commit()
        db.refresh(new_song)

        # Aktualizujeme dokumenty
        sync_ma_documents(email, drv, db)

        return new_song
    except HTTPException:
        raise
    except Exception as e:
        print(f"MA Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/songs/{song_id}")
def delete_song(song_id: int, request: Request, db: Session = Depends(get_db)):
    email = get_user_email(request)
    song = db.query(Song).filter(Song.id == song_id, Song.owner_email == email).first()
    if not song:
        raise HTTPException(status_code=404, detail="Skladba nenalezena")

    db.delete(song)
    db.commit()

    # Synchronizace dokumentů
    try:
        sid = request.session.get("sid")
        token_data = TOKENS[sid]
        creds = get_credentials_from_session(token_data["token"])
        drv = drive_service(creds)
        sync_ma_documents(email, drv, db)
    except:
        pass

    return {"status": "ok"}


@router.patch("/songs/{song_id}", response_model=SongOut)
def update_song(
    song_id: int, payload: SongUpdate, request: Request, db: Session = Depends(get_db)
):
    email = get_user_email(request)
    song = db.query(Song).filter(Song.id == song_id, Song.owner_email == email).first()
    if not song:
        raise HTTPException(status_code=404, detail="Skladba nenalezena")

    if payload.number is not None:
        num = payload.number.strip()
        if num.upper() == "N":
            num = "N"
        elif not num.isdigit():
            raise HTTPException(
                status_code=400, detail="Číslo písně musí být celé číslo nebo 'N'"
            )

        if num != "N" and num != song.number:
            existing = (
                db.query(Song)
                .filter(Song.owner_email == email, Song.number == num)
                .first()
            )
            if existing:
                raise HTTPException(
                    status_code=400, detail=f"Skladba s číslem {num} již existuje"
                )

        song.number = num

    if payload.title is not None:
        song.title = payload.title
    if payload.singer is not None:
        song.singer = payload.singer
    if payload.duration is not None:
        song.duration = payload.duration
    if payload.category is not None:
        song.category = payload.category

    # Přejmenování složky na Disku, pokud se změnilo číslo nebo název
    if (
        payload.number is not None or payload.title is not None
    ) and song.drive_folder_id:
        try:
            sid = request.session.get("sid")
            token_data = TOKENS[sid]
            creds = get_credentials_from_session(token_data["token"])
            drv = drive_service(creds)

            new_name = f"{song.number} {song.title}"
            drv.files().update(
                fileId=song.drive_folder_id, body={"name": new_name}
            ).execute()
        except Exception as e:
            print(f"Drive rename error: {e}")

    db.commit()
    db.refresh(song)

    # Synchronizace dokumentů
    try:
        sid = request.session.get("sid")
        token_data = TOKENS[sid]
        creds = get_credentials_from_session(token_data["token"])
        drv = drive_service(creds)
        sync_ma_documents(email, drv, db)
    except:
        pass

    return song


from app.services.pdf_splitter import segment_pdf, extract_pdf_pages, match_instrument_name
from pypdf import PdfReader
import json


@router.post("/songs/{song_id}/analyze_pdf")
async def analyze_song_pdf(
    song_id: int,
    request: Request,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    email = get_user_email(request)
    song = (
        db.query(Song)
        .options(joinedload(Song.files))
        .filter(Song.id == song_id, Song.owner_email == email)
        .first()
    )
    if not song:
        raise HTTPException(status_code=404, detail="Skladba nenalezena")

    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Soubor musí být ve formátu PDF.")

    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Nahraný soubor je prázdný.")

    # Získat nástrojový setup kapely
    instruments = (
        db.query(Instrument)
        .filter(Instrument.owner_email == email)
        .all()
    )
    inst_dicts = [
        {"id": inst.id, "name": inst.name, "category": inst.category, "is_tracked": inst.is_tracked}
        for inst in instruments
    ]

    existing_files_dicts = [
        {
            "id": f.id,
            "name": f.name,
            "file_type": f.file_type,
            "instrument_name": f.instrument_name,
        }
        for f in song.files
    ]

    segments = segment_pdf(content, song.title, inst_dicts, existing_files_dicts)

    total_pages = 0
    try:
        reader = PdfReader(io.BytesIO(content))
        total_pages = len(reader.pages)
    except:
        pass

    return {
        "filename": file.filename,
        "total_pages": total_pages,
        "segments": segments,
    }


@router.post("/songs/{song_id}/analyze_files")
async def analyze_song_files(
    song_id: int,
    request: Request,
    files: list[UploadFile] = File(...),
    db: Session = Depends(get_db),
):
    email = get_user_email(request)
    song = db.query(Song).filter(Song.id == song_id, Song.owner_email == email).first()
    if not song:
        raise HTTPException(status_code=404, detail="Skladba nenalezena")

    instruments = (
        db.query(Instrument)
        .filter(Instrument.owner_email == email)
        .all()
    )
    all_band_names = [inst.name for inst in instruments]

    results = []
    for file in files:
        filename = file.filename or "unknown"
        ext = os.path.splitext(filename)[1].lower()

        detected_type = "other"
        detected_inst = None

        if ext in [".mp3", ".wav", ".mid", ".midi", ".m4a"]:
            results.append({
                "filename": filename,
                "file_type": "audio",
                "instrument_name": None,
            })
            continue

        if ext == ".pdf":
            try:
                content = await file.read()
                if len(content) > 0:
                    reader = PdfReader(io.BytesIO(content))
                    if len(reader.pages) > 0:
                        first_page = reader.pages[0]
                        raw_text = first_page.extract_text() or ""
                        lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
                        header_text = " | ".join(lines[:4]) if lines else ""

                        # Zkusíme záhlaví PDF
                        ft, inst_name, conf = match_instrument_name(
                            header_text, raw_text, all_band_names
                        )
                        if ft != "other" or inst_name:
                            detected_type = ft
                            detected_inst = inst_name
            except Exception as e:
                print(f"Analyze multi-file PDF error for {filename}: {e}")

        # Pokud se ze záhlaví nepodařilo určit nebo není PDF, zkusíme název souboru
        if detected_type == "other" and not detected_inst:
            ft, inst_name, conf = match_instrument_name(
                filename, filename, all_band_names
            )
            if ft != "other" or inst_name:
                detected_type = ft
                detected_inst = inst_name

        results.append({
            "filename": filename,
            "file_type": detected_type,
            "instrument_name": detected_inst,
        })

    return results


@router.post("/songs/{song_id}/split_and_upload")
async def split_and_upload_song_pdf(
    song_id: int,
    request: Request,
    file: UploadFile = File(...),
    rules_json: str = Form(...),
    db: Session = Depends(get_db),
):
    email = get_user_email(request)
    song = (
        db.query(Song)
        .options(joinedload(Song.files))
        .filter(Song.id == song_id, Song.owner_email == email)
        .first()
    )
    if not song:
        raise HTTPException(status_code=404, detail="Skladba nenalezena")

    if not song.drive_folder_id:
        raise HTTPException(status_code=400, detail="Skladba nemá složku na Disku")

    try:
        rules = json.loads(rules_json)
        if not isinstance(rules, list) or len(rules) == 0:
            raise HTTPException(status_code=400, detail="Nebyly vybrány žádné party k uložení.")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Neplatná pravidla rozdělení: {e}")

    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Nahraný soubor je prázdný.")

    sid = request.session.get("sid")
    token_data = TOKENS[sid]
    creds = get_credentials_from_session(token_data["token"])
    drv = drive_service(creds)

    dash = " – "
    uploaded_results = []

    for rule in rules:
        file_type = rule.get("file_type", "part")
        instrument_name = rule.get("instrument_name")
        page_start = int(rule.get("page_start", 1))
        page_end = int(rule.get("page_end", 1))

        part_pdf_bytes = extract_pdf_pages(content, page_start, page_end)
        file_like = io.BytesIO(part_pdf_bytes)

        if file_type == "part":
            display_name = instrument_name if instrument_name else f"Part (str. {page_start}-{page_end})"
            new_filename = f"{display_name}{dash}{song.title}.pdf"
        elif file_type == "score":
            new_filename = f"Partitura{dash}{song.title}.pdf"
        else:
            new_filename = f"Jiné (str. {page_start}-{page_end}){dash}{song.title}.pdf"

        drive_file = upload_file_to_drive(
            drv, song.drive_folder_id, file_like, new_filename, "application/pdf"
        )

        existing_file = None
        if file_type == "part":
            existing_file = (
                db.query(SongFile)
                .filter(
                    SongFile.song_id == song_id,
                    SongFile.file_type == "part",
                    SongFile.instrument_name == instrument_name,
                )
                .first()
            )
        elif file_type == "score":
            existing_file = (
                db.query(SongFile)
                .filter(SongFile.song_id == song_id, SongFile.file_type == "score")
                .first()
            )

        if existing_file:
            existing_file.drive_file_id = drive_file["id"]
            existing_file.name = new_filename
        else:
            new_song_file = SongFile(
                song_id=song_id,
                drive_file_id=drive_file["id"],
                name=new_filename,
                file_type=file_type,
                instrument_name=instrument_name if file_type == "part" else None,
            )
            db.add(new_song_file)

        uploaded_results.append({
            "name": new_filename,
            "instrument_name": instrument_name,
            "drive_file_id": drive_file["id"],
        })

    db.commit()

    # Synchronizace Word dokumentu
    sync_ma_documents(email, drv, db)

    return {
        "status": "ok",
        "uploaded_count": len(uploaded_results),
        "files": uploaded_results,
    }


@router.post("/songs/{song_id}/files")
async def upload_song_file(
    song_id: int,
    request: Request,
    file: UploadFile = File(...),
    file_type: str = Form(...),  # 'part' | 'score' | 'audio' | 'other'
    instrument_name: Optional[str] = Form(None),
    db: Session = Depends(get_db),
):
    email = get_user_email(request)
    song = db.query(Song).filter(Song.id == song_id, Song.owner_email == email).first()
    if not song:
        raise HTTPException(status_code=404, detail="Skladba nenalezena")

    if not song.drive_folder_id:
        raise HTTPException(status_code=400, detail="Skladba nemá složku na Disku")

    try:
        # Načtení obsahu do paměti pro maximální kompatibilitu s Google API
        content = await file.read()
        file_like = io.BytesIO(content)

        # Google Drive Logika
        sid = request.session.get("sid")
        token_data = TOKENS[sid]
        creds = get_credentials_from_session(token_data["token"])
        drv = drive_service(creds)

        # Pojmenovávací konvence: [Nástroj/Partitura] – [Název skladby].[ext]
        ext = os.path.splitext(file.filename)[1]
        dash = " – "  # Návrat k hezké dlouhé pomlčce

        if file_type == "part":
            display_name = instrument_name if instrument_name else "Neznámý nástroj"
            new_filename = f"{display_name}{dash}{song.title}{ext}"
        elif file_type == "score":
            new_filename = f"Partitura{dash}{song.title}{ext}"
        else:
            new_filename = file.filename

        # Upload
        drive_file = upload_file_to_drive(
            drv, song.drive_folder_id, file_like, new_filename, file.content_type
        )

        # Uložit do DB
        existing_file = None
        if file_type == "part":
            existing_file = (
                db.query(SongFile)
                .filter(
                    SongFile.song_id == song_id,
                    SongFile.file_type == "part",
                    SongFile.instrument_name == instrument_name,
                )
                .first()
            )
        elif file_type == "score":
            existing_file = (
                db.query(SongFile)
                .filter(SongFile.song_id == song_id, SongFile.file_type == "score")
                .first()
            )

        if existing_file:
            existing_file.drive_file_id = drive_file["id"]
            existing_file.name = new_filename
        else:
            new_song_file = SongFile(
                song_id=song_id,
                drive_file_id=drive_file["id"],
                name=new_filename,
                file_type=file_type,
                instrument_name=instrument_name if file_type == "part" else None,
            )
            db.add(new_song_file)

        db.commit()

        # Aktualizujeme dokumenty
        sync_ma_documents(email, drv, db)

        return {"status": "ok", "file_id": drive_file["id"]}
    except Exception as e:
        print(f"MA Upload Detail Error: {str(e)}")
        import traceback

        traceback.print_exc()
        raise HTTPException(
            status_code=500, detail=f"Chyba serveru při nahrávání: {str(e)}"
        )


def sync_ma_documents(email: str, drv, db: Session):
    try:
        songs = (
            db.query(Song)
            .options(joinedload(Song.files))
            .filter(Song.owner_email == email)
            .all()
        )

        def sort_key(s):
            cat_prio = 0 if s.category in ("Standard", "Standardní repertoár") else 1
            num_is_n = 1 if s.number == "N" else 0
            try:
                num_val = int(s.number)
            except:
                num_val = 0
            return (cat_prio, num_is_n, num_val, s.title)

        songs.sort(key=sort_key)

        # 1. Seznam skladeb

        doc = Document()

        # Nastavení stránky A4 s menšími okraji
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # Základní styl – Calibri Bold 18
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(18)

        # Nadpis
        heading = doc.add_paragraph()
        heading.alignment = 1  # CENTER
        hr = heading.add_run("SEZNAM SKLADEB")
        hr.font.size = Pt(24)
        hr.font.bold = True
        heading.paragraph_format.space_after = Pt(18)

        # Každý řádek: číslo+název [TAB] zpěvák
        # A4 usable width ≈ 17 cm = ~483 pt; tab stop pro zpěváka na 340 pt (~12 cm od okraje)
        SINGER_TAB = Pt(340)

        # Přejmenování kategorie pro zobrazení
        def cat_display(cat):
            if cat in ("Standard", "Standardní repertoár"):
                return "Standardní repertoár"
            return cat

        # Seskupení podle kategorií
        from itertools import groupby

        for cat_raw, group in groupby(songs, key=lambda s: s.category):
            cat_label = cat_display(cat_raw)

            # Podnadpis kategorie
            p_cat = doc.add_paragraph()
            p_cat.paragraph_format.space_before = Pt(12)
            p_cat.paragraph_format.space_after = Pt(4)
            r_cat = p_cat.add_run(f"{cat_label}:")
            r_cat.font.name = "Calibri"
            r_cat.font.bold = True
            r_cat.font.size = Pt(14)

            for s in group:
                p = doc.add_paragraph()
                p.paragraph_format.tab_stops.add_tab_stop(
                    SINGER_TAB, WD_TAB_ALIGNMENT.LEFT
                )
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)

                # Číslo N → jen název bez čísla
                if s.number == "N":
                    title_str = s.title.upper()
                else:
                    title_str = f"{s.number}. {s.title.upper()}"

                r_title = p.add_run(title_str)
                r_title.font.name = "Calibri"
                r_title.font.bold = True
                r_title.font.size = Pt(18)
                p.add_run("\t")
                r_singer = p.add_run(s.singer.upper())
                r_singer.font.name = "Calibri"
                r_singer.font.bold = True
                r_singer.font.size = Pt(18)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        root_id = os.getenv("BAND_DRIVE_ROOT_FOLDER_ID")
        update_or_create_file(
            drv,
            root_id,
            buf,
            "Aktuální seznam skladeb.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        print(f"MA Sync Doc Error: {e}")


@router.post("/documents/generate")
def generate_ma_documents(request: Request, db: Session = Depends(get_db)):
    """Vygeneruje (nebo aktualizuje) Word dokument Aktuální seznam skladeb na Drive a vrátí jeho přímý odkaz."""
    try:
        email = get_user_email(request)
        sid = request.session.get("sid")
        token_data = TOKENS[sid]
        creds = get_credentials_from_session(token_data["token"])
        drv = drive_service(creds)

        # Vygeneruj / aktualizuj dokument
        sync_ma_documents(email, drv, db)

        root_id = os.getenv("BAND_DRIVE_ROOT_FOLDER_ID")

        from app.google_client import find_file_id

        list_id = find_file_id(drv, root_id, "Aktuální seznam skladeb.docx")

        def get_link(fid):
            if not fid:
                return None
            f = drv.files().get(fileId=fid, fields="webViewLink").execute()
            return f.get("webViewLink")

        return {
            "current_list": get_link(list_id),
        }
    except Exception as e:
        print(f"MA generate docs error: {e}")
        import traceback

        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/playlist/parse_pdf")
async def parse_playlist_pdf_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    try:
        pdf_bytes = await file.read()
        from app.routers.events import parse_pdf_to_playlist_structure
        parsed = parse_pdf_to_playlist_structure(pdf_bytes, db)
        return parsed
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Chyba při parsování PDF: {str(e)}"
        )
